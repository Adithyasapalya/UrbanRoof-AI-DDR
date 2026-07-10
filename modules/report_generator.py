"""
==========================================================
UrbanRoof AI DDR Generator

Professional Report Generator

Author : Adithya Sapalya
==========================================================
"""

from __future__ import annotations

import os
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from modules.knowledge_base import KnowledgeBase


class ReportGenerator:

    def __init__(self):

        self.document = Document()

        self.document.core_properties.author = "Adithya Sapalya"
        self.document.core_properties.title = "UrbanRoof AI DDR Report"
        self.document.core_properties.subject = "Defect Diagnostic Report"

        self._create_styles()

    # --------------------------------------------------

    def safe_text(self, value, default="Not Available"):

        if value is None:
            return default

        return str(value)

    # --------------------------------------------------

    def _create_styles(self):

        styles = self.document.styles

        if "HeadingBlue" not in styles:

            style = styles.add_style(
                "HeadingBlue",
                WD_STYLE_TYPE.PARAGRAPH
            )

            style.font.name = "Calibri"
            style.font.size = Pt(16)
            style.font.bold = True

        if "BodyTextCustom" not in styles:

            style = styles.add_style(
                "BodyTextCustom",
                WD_STYLE_TYPE.PARAGRAPH
            )

            style.font.name = "Calibri"
            style.font.size = Pt(11)

    # --------------------------------------------------

    def title_page(self):

        title = self.document.add_heading(
            "UrbanRoof AI\nDetailed Diagnostic Report",
            level=0
        )
        title = self.document.add_paragraph(
            "Made Using Python and Artificial Intelligence" \
            "\n\nAuthor: Adithya Sapalya" \
            "\n\nDate: 2024-06-15" \
            "\n\nDisclaimer: This report is generated using AI and should be verified by a qualified professional." \
            "n\nNote: The information provided in this report is based on the analysis of the inspection and thermal reports. It is intended for informational purposes only and should not be considered as a substitute for professional advice or judgment." \
            "\n\nThe author and UrbanRoof AI are not responsible for any decisions made based on the information contained in this report. Users are advised to consult with qualified professionals for any actions or decisions related to the property." \
            
        )

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run(
            "Generated using Artificial Intelligence\n"
        ).bold = True

        p.add_run(
            "Inspection + Thermal Report Analysis"
        )

        self.document.add_page_break()

    # --------------------------------------------------

    def executive_summary(self, llm_report):

        self.document.add_heading(
            "1. Executive Summary",
            level=1
        )

        summary = llm_report.get(
            "executive_summary",
            {}
        )

        self.document.add_paragraph(

            summary.get(
                "executive_summary",
                "Not Available"
            )

        )

        self.document.add_paragraph()

        self.document.add_paragraph(

            "Overall Condition : "
            + summary.get(
                "overall_condition",
                "Unknown"
            )

        )

        actions = summary.get(
            "priority_actions",
            []
        )

        if actions:

            self.document.add_heading(
                "Priority Actions",
                level=2
            )

            for action in actions:

                self.document.add_paragraph(
                    action,
                    style="List Bullet"
                )

    # --------------------------------------------------

    def property_issue_summary(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "2. Property Issue Summary",
            level=1
        )

        issue_count = defaultdict(int)

        for obs in kb.get_all_observations():

            issue = obs.issue if obs.issue else "Unknown Issue"

            issue_count[issue] += 1

        table = self.document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        hdr = table.rows[0].cells

        hdr[0].text = "Issue"
        hdr[1].text = "Occurrences"

        for issue, count in sorted(issue_count.items()):

            row = table.add_row().cells

            row[0].text = issue
            row[1].text = str(count)
        # -----------------------------------------------------
    # Property Issue Summary
    # -----------------------------------------------------

    def property_issue_summary(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "2. Property Issue Summary",
            level=1
        )

        issue_count = defaultdict(int)
        severity_count = defaultdict(int)

        for obs in kb.get_all_observations():

            issue = obs.issue if obs.issue else "Unknown Issue"
            severity = obs.severity if obs.severity else "Unknown"

            issue_count[issue] += 1
            severity_count[severity] += 1

        self.document.add_heading(
            "Issue Distribution",
            level=2
        )

        table = self.document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "Issue"
        hdr[1].text = "Occurrences"

        for issue, count in sorted(issue_count.items()):

            row = table.add_row().cells
            row[0].text = issue
            row[1].text = str(count)

        self.document.add_paragraph()

        self.document.add_heading(
            "Severity Distribution",
            level=2
        )

        sev_table = self.document.add_table(
            rows=1,
            cols=2
        )

        sev_table.style = "Table Grid"

        hdr = sev_table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"

        for sev in [
            "Critical",
            "High",
            "Medium",
            "Low",
            "Unknown"
        ]:

            row = sev_table.add_row().cells
            row[0].text = sev
            row[1].text = str(
                severity_count.get(sev, 0)
            )

        self.document.add_page_break()

        # -----------------------------------------------------
    # Area-wise Observations
    # -----------------------------------------------------

    def area_wise_observations(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "3. Area-wise Observations",
            level=1
        )

        grouped = defaultdict(list)

        for obs in kb.get_all_observations():
            grouped[obs.area].append(obs)

        for area in sorted(grouped.keys()):

            self.document.add_heading(
                area,
                level=2
            )

            observations = grouped[area]

            for index, obs in enumerate(observations, start=1):

                # -----------------------------
                # Auto Severity
                # -----------------------------

                text = (
                    (obs.description or "") +
                    " " +
                    (obs.issue or "")
                ).lower()

                if any(word in text for word in [
                    "collapse",
                    "major crack",
                    "structural",
                    "unsafe",
                    "reinforcement exposed",
                    "critical"
                ]):

                    obs.severity = "Critical"

                elif any(word in text for word in [
                    "crack",
                    "leak",
                    "water",
                    "seepage",
                    "rust",
                    "corrosion",
                    "damaged"
                ]):

                    obs.severity = "High"

                elif any(word in text for word in [
                    "moisture",
                    "thermal",
                    "vegetation",
                    "efflorescence",
                    "hollow",
                    "damp"
                ]):

                    obs.severity = "Medium"

                else:

                    obs.severity = "Low"

                # -----------------------------
                # Observation Heading
                # -----------------------------

                self.document.add_heading(
                    f"Observation {index}",
                    level=3
                )

                table = self.document.add_table(
                    rows=8,
                    cols=2
                )

                table.style = "Table Grid"

                table.cell(0,0).text = "Issue"
                table.cell(0,1).text = self.safe_text(obs.issue)

                table.cell(1,0).text = "Description"
                table.cell(1,1).text = self.safe_text(obs.description)

                table.cell(2,0).text = "Severity"
                table.cell(2,1).text = self.safe_text(obs.severity)

                table.cell(3,0).text = "Root Cause"
                table.cell(3,1).text = (
                    obs.root_cause
                    if obs.root_cause
                    else "Not Available"
                )

                table.cell(4,0).text = "Recommendation"
                table.cell(4,1).text = (
                    obs.recommendation
                    if obs.recommendation
                    else "Not Available"
                )

                table.cell(5,0).text = "Similarity Score"

                table.cell(5,1).text = (
                    f"{obs.similarity_score:.2f}"
                    if obs.similarity_score
                    else "Not Matched"
                )

                table.cell(6,0).text = "Matched Observation"

                table.cell(6,1).text = (
                    str(obs.matched_observation_id)
                    if obs.matched_observation_id is not None
                    else "Not Available"
                )

                table.cell(7,0).text = "Evidence"

                table.cell(7,1).text = self.safe_text(
                    obs.source_evidence)
                
                # -----------------------------
                # Inspection Image
                # -----------------------------

                self.document.add_paragraph()

                self.document.add_heading(
                    "Inspection Image",
                    level=4
                )

                if obs.image_refs:

                    inserted = False

                    for img in obs.image_refs:

                        if img and os.path.exists(img):

                            try:

                                self.document.add_picture(
                                    img,
                                    width=Inches(5.5)
                                )

                                inserted = True

                            except Exception as e:

                                self.document.add_paragraph(
                                    f"Could not load image:\n{e}"
                                )

                    if not inserted:

                        self.document.add_paragraph(
                            "Image Not Available"
                        )

                else:

                    self.document.add_paragraph(
                        "Image Not Available"
                    )

                self.document.add_page_break()
                

        # -----------------------------------------------------
    # Severity Assessment
    # -----------------------------------------------------

    def severity_assessment(
        self,
        kb: KnowledgeBase,
        llm_report
    ):

        self.document.add_heading(
            "4. Severity Assessment",
            level=1
        )

        severity_stats = {
            "Critical": 0,
            "High": 0,
            "Medium": 0,
            "Low": 0,
            "Unknown": 0
        }

        total = 0

        for obs in kb.get_all_observations():

            sev = (obs.severity or "Unknown").title()

            if sev not in severity_stats:
                sev = "Unknown"

            severity_stats[sev] += 1
            total += 1

        table = self.document.add_table(
            rows=1,
            cols=3
        )

        table.style = "Table Grid"

        hdr = table.rows[0].cells

        hdr[0].text = "Severity"
        hdr[1].text = "Count"
        hdr[2].text = "Percentage"

        for sev in [
            "Critical",
            "High",
            "Medium",
            "Low",
            "Unknown"
        ]:

            row = table.add_row().cells

            count = severity_stats[sev]

            percentage = (
                (count / total) * 100
                if total > 0 else 0
            )

            row[0].text = sev
            row[1].text = str(count)
            row[2].text = f"{percentage:.1f}%"

        self.document.add_paragraph()

        if severity_stats["Critical"] > 0:

            summary = (
                "Critical structural issues were detected. "
                "Immediate engineering intervention is recommended."
            )

        elif severity_stats["High"] > 0:

            summary = (
                "Several high severity defects were identified. "
                "Repair should be prioritized."
            )

        elif severity_stats["Medium"] > 0:

            summary = (
                "Moderate deterioration was identified. "
                "Planned maintenance is recommended."
            )

        else:

            summary = (
                "Only minor observations were identified. "
                "Routine monitoring is sufficient."
            )

        self.document.add_paragraph(summary)

        self.document.add_page_break()


    # -----------------------------------------------------
    # Recommended Actions
    # -----------------------------------------------------

    def recommended_actions(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "5. Recommended Actions",
            level=1
        )

        immediate = []
        planned = []
        monitoring = []

        for obs in kb.get_all_observations():

            recommendation = (
                obs.recommendation.strip()
                if obs.recommendation
                else ""
            )

            if recommendation == "":
                recommendation = (
                    f"Inspect and repair {obs.issue.lower()} in {obs.area}."
                )

            severity = (
                obs.severity.lower()
                if obs.severity
                else "low"
            )

            if severity in [
                "critical",
                "high"
            ]:

                immediate.append(recommendation)

            elif severity == "medium":

                planned.append(recommendation)

            else:

                monitoring.append(recommendation)

        # ----------------------------

        self.document.add_heading(
            "Immediate Actions",
            level=2
        )

        if immediate:

            for action in sorted(set(immediate)):
                self.document.add_paragraph(
                    action,
                    style="List Bullet"
                )

        else:

            self.document.add_paragraph(
                "No immediate actions required."
            )

        # ----------------------------

        self.document.add_heading(
            "Planned Maintenance",
            level=2
        )

        if planned:

            for action in sorted(set(planned)):
                self.document.add_paragraph(
                    action,
                    style="List Bullet"
                )

        else:

            self.document.add_paragraph(
                "No planned maintenance required."
            )

        # ----------------------------

        self.document.add_heading(
            "Routine Monitoring",
            level=2
        )

        if monitoring:

            for action in sorted(set(monitoring)):
                self.document.add_paragraph(
                    action,
                    style="List Bullet"
                )

        else:

            self.document.add_paragraph(
                "Routine inspection only."
            )

        self.document.add_page_break()

        # -----------------------------------------------------
    # Additional Notes
    # -----------------------------------------------------

    def additional_notes(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "6. Additional Notes",
            level=1
        )

        notes = [

            "Inspection observations were correlated with thermal evidence where available.",

            "Severity has been assigned using rule-based AI analysis.",

            "Recommendations are advisory and should be verified by a qualified structural engineer.",

            "Thermal anomalies do not always indicate structural defects and require engineering judgement.",

            f"Total observations analysed : {len(kb.get_all_observations())}"

        ]

        for note in notes:
            self.document.add_paragraph(
                note,
                style="List Bullet"
            )

        self.document.add_page_break()

    # -----------------------------------------------------
    # Missing Information
    # -----------------------------------------------------

    def missing_information(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "7. Missing Information",
            level=1
        )

        missing = False

        for obs in kb.get_all_observations():

            if not obs.root_cause:

                self.document.add_paragraph(
                    f"{obs.area}: Root cause unavailable.",
                    style="List Bullet"
                )

                missing = True

            if not obs.recommendation:

                self.document.add_paragraph(
                    f"{obs.area}: Recommendation unavailable.",
                    style="List Bullet"
                )

                missing = True

        if not missing:

            self.document.add_paragraph(
                "No missing information detected."
            )

        self.document.add_page_break()

    # -----------------------------------------------------
    # Building Health Score
    # -----------------------------------------------------

    def building_health(
        self,
        llm_report
    ):

        self.document.add_heading(
            "8. Building Health Score",
            level=1
        )

        score = llm_report.get(
            "health_score",
            75
        )

        p = self.document.add_paragraph()

        p.add_run(
            "Building Health Score : "
        ).bold = True

        p.add_run(
            f"{score}/100"
        )

        if score >= 90:
            status = "Excellent"

        elif score >= 75:
            status = "Good"

        elif score >= 60:
            status = "Fair"

        elif score >= 40:
            status = "Poor"

        else:
            status = "Critical"

        self.document.add_paragraph(
            f"Overall Building Condition : {status}"
        )

        self.document.add_page_break()

    # -----------------------------------------------------
    # Appendix
    # -----------------------------------------------------

    def appendix(
        self,
        kb: KnowledgeBase
    ):

        self.document.add_heading(
            "9. Appendix",
            level=1
        )

        table = self.document.add_table(
            rows=1,
            cols=5
        )

        table.style = "Table Grid"

        hdr = table.rows[0].cells

        hdr[0].text = "ID"
        hdr[1].text = "Area"
        hdr[2].text = "Issue"
        hdr[3].text = "Severity"
        hdr[4].text = "Matched"

        for obs in kb.get_all_observations():

            row = table.add_row().cells

            row[0].text = str(obs.id)

            row[1].text = self.safe_text(obs.area)

            row[2].text = self.safe_text(obs.issue)

            row[3].text = self.safe_text(obs.severity)

            row[4].text = (
                str(obs.matched_observation_id)
                if obs.matched_observation_id is not None
                else "-"
            )

    # -----------------------------------------------------
    # Generate Report
    # -----------------------------------------------------

    def generate_report(
        self,
        kb,
        llm_report
    ):

        print("\nGenerating DDR Report...")

        self.title_page()

        self.executive_summary(llm_report)

        self.property_issue_summary(kb)

        self.area_wise_observations(kb)

        self.severity_assessment(
            kb,
            llm_report
        )

        self.recommended_actions(kb)

        self.additional_notes(kb)

        self.missing_information(kb)

        self.building_health(llm_report)

        self.appendix(kb)

        print("Report Generated Successfully.")

        return self.document

    # -----------------------------------------------------
    # Save Report
    # -----------------------------------------------------

    def save(
        self,
        output_file
    ):

        output_path = os.path.abspath(output_file)

        os.makedirs(
            os.path.dirname(output_path),
            exist_ok=True
        )

        try:

            self.document.save(output_path)

            print(f"Report saved -> {output_path}")

        except PermissionError:

            backup = output_path.replace(
                ".docx",
                "_new.docx"
            )

            self.document.save(backup)

            print(f"Report saved -> {backup}")

    # -----------------------------------------------------
    # Run
    # -----------------------------------------------------

    def run(
        self,
        kb,
        llm_report,
        output_file
    ):

        self.generate_report(
            kb,
            llm_report
        )

        self.save(output_file)

        return self.document    